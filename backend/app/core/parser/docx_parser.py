from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table
from typing import List, Union
import os


class DocxParser:
    """Word document parser for loading and accessing document structure"""

    def __init__(self, file_path: str):
        """
        Initialize parser with Word document path

        Args:
            file_path: Path to .docx file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.endswith('.docx'):
            raise ValueError("Only .docx files are supported")

        self.file_path = file_path
        self.document: DocumentType = Document(file_path)

    def get_paragraphs(self) -> List[Paragraph]:
        """Get all paragraphs from document"""
        return self.document.paragraphs

    def get_tables(self) -> List[Table]:
        """Get all tables from document"""
        return self.document.tables

    def get_body_elements(self) -> List[Union[Paragraph, Table]]:
        """
        Get all body elements (paragraphs and tables) in document order

        Returns:
            List of paragraphs and tables in the order they appear
        """
        elements = []
        for element in self.document.element.body:
            if element.tag.endswith('p'):
                # Paragraph element
                for para in self.document.paragraphs:
                    if para._element == element:
                        elements.append(para)
                        break
            elif element.tag.endswith('tbl'):
                # Table element
                for table in self.document.tables:
                    if table._element == element:
                        elements.append(table)
                        break
        return elements

    def get_paragraph_text(self, paragraph: Paragraph) -> str:
        """Get plain text from paragraph"""
        return paragraph.text

    def get_paragraph_runs(self, paragraph: Paragraph):
        """Get all runs (text segments with formatting) from paragraph"""
        return paragraph.runs

    def has_images(self) -> bool:
        """Check if document contains images"""
        for rel in self.document.part.rels.values():
            if "image" in rel.target_ref:
                return True
        return False

    def get_image_parts(self):
        """Get all image parts from document"""
        image_parts = []
        for rel in self.document.part.rels.values():
            if "image" in rel.target_ref:
                image_parts.append(rel.target_part)
        return image_parts
