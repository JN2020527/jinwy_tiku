import re
from typing import List, Dict, Optional, Tuple
from docx.text.paragraph import Paragraph


class QuestionSection:
    """Represents a question type section"""

    def __init__(self, type_name: str, start_index: int):
        self.type_name = type_name
        self.start_index = start_index
        self.questions: List[QuestionBlock] = []


class QuestionBlock:
    """Represents a single question block"""

    def __init__(self, number: str, start_index: int):
        self.number = number
        self.start_index = start_index
        self.end_index: Optional[int] = None
        self.paragraphs: List[Paragraph] = []
        self.is_material_question = False
        self.sub_questions: List["QuestionBlock"] = []


class StructureParser:
    """Parser for recognizing document structure (question types and numbers)"""

    CHINESE_NUMERALS = {
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }

    TYPE_PATTERN = re.compile(r"^([一二三四五六七八九十]+)、\s*(.+)$")
    QUESTION_NUMBER_PATTERN = re.compile(r"^(\d+)[．.]\s*")
    SUB_QUESTION_PATTERN = re.compile(r"^\((\d+)\)\s*")
    ANALYSIS_DETAIL_PATTERN = re.compile(r"^(\d+)[．.]\s*([A-D])[、．.]")
    ANALYSIS_OPTION_DETAIL_PATTERN = re.compile(r"^([A-D])、")

    MATERIAL_KEYWORDS = [
        r"阅读下列材料，完成下面小题",
        r"阅读下列材料，回答下列问题",
        r"阅读下列材料，回答问题",
        r"回答下列小题",
        r"完成下列题目",
        r"完成下面小题",  # Added missing keyword
    ]

    def __init__(self, paragraphs: List[Paragraph]):
        """Initialize structure parser

        Args:
            paragraphs: List of paragraphs from document
        """
        self.paragraphs = paragraphs
        self.sections: List[QuestionSection] = []

    def _paragraph_has_image(self, paragraph: Paragraph) -> bool:
        """Check if paragraph contains embedded images."""
        for run in paragraph.runs:
            try:
                if run._element.xpath('.//a:blip') or run._element.xpath('.//v:imagedata'):
                    return True
            except Exception:
                continue
        return False

    def parse(self) -> List[QuestionSection]:
        """Parse document structure to identify question sections and blocks

        Returns:
            List of question sections with their questions
        """
        current_section: Optional[QuestionSection] = None
        current_question: Optional[QuestionBlock] = None
        in_material = False
        has_material_content = False

        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            has_image = self._paragraph_has_image(para)
            is_fill_section = current_section and "填空" in current_section.type_name
            is_inline_sub_section = current_section and any(
                key in current_section.type_name
                for key in ["填空", "简答", "实验", "解答", "综合应用"]
            )

            if not text and not has_image:
                continue

            # Check for question type section
            type_match = self.TYPE_PATTERN.match(text)
            if type_match:
                chinese_num, type_name = type_match.groups()
                current_section = QuestionSection(type_name.strip(), i)
                self.sections.append(current_section)
                current_question = None
                in_material = False
                has_material_content = False
                continue

            # Check for material question keywords
            if not is_fill_section and any(keyword in text.strip() for keyword in self.MATERIAL_KEYWORDS):
                # Close previous question before starting new material collection
                if current_question:
                    current_question.end_index = i - 1

                in_material = True
                # Create a placeholder question to hold material content
                current_question = QuestionBlock("material_placeholder", i)
                if current_section:
                    current_section.questions.append(current_question)
                current_question.paragraphs.append(para)
                has_material_content = True
                continue

            # Collect material content paragraphs (only before sub-questions are created)
            if in_material and has_material_content and current_question and not current_question.sub_questions and not self.QUESTION_NUMBER_PATTERN.match(text) and not self.SUB_QUESTION_PATTERN.match(text):
                # Check if this is an attribute block (答案/难度/知识点/解析/详解)
                # Always add attribute paragraphs to parent (material) question, not to sub-questions
                # This ensures attributes are available for parsing but don't appear in sub-question stems
                if any(attr in text for attr in ["【答案】", "【难度】", "【知识点】", "【解析】", "【详解】"]):
                    current_question.paragraphs.append(para)
                    continue

                # Add material content to the placeholder question (only before sub-questions exist)
                current_question.paragraphs.append(para)
                continue

            # After sub-questions are created, attribute blocks should still go to parent
            if in_material and has_material_content and current_question and current_question.sub_questions:
                if any(attr in text for attr in ["【答案】", "【难度】", "【知识点】", "【解析】", "【详解】"]):
                    current_question.paragraphs.append(para)
                    continue

            # Check if this is analysis detail paragraph (e.g., "4．A、解析内容")
            # These should be added to parent question, not treated as sub-questions
            if in_material and has_material_content and current_question and self.ANALYSIS_DETAIL_PATTERN.match(text):
                current_question.paragraphs.append(para)
                continue

            # Check if this is analysis option detail (e.g., "A、解析内容")
            # These appear in 【解析】block and should be added to parent
            if in_material and has_material_content and current_question and current_question.sub_questions and self.ANALYSIS_OPTION_DETAIL_PATTERN.match(text):
                current_question.paragraphs.append(para)
                continue

            question_match = self.QUESTION_NUMBER_PATTERN.match(text)
            if question_match:
                number = question_match.group(1)

                # If in material phase, this is a sub-question (either first or subsequent)
                if in_material and has_material_content and current_question:
                    # If this is the first sub-question, update placeholder number
                    if current_question.number == "material_placeholder":
                        current_question.number = number
                        current_question.is_material_question = True

                    # Create sub-question with question number
                    sub_question = QuestionBlock(number, i)
                    sub_question.paragraphs.append(para)
                    current_question.sub_questions.append(sub_question)
                    continue

                # Close previous question (normal case)
                if current_question and not in_material:
                    current_question.end_index = i - 1

                # Create new question (normal case)
                current_question = QuestionBlock(number, i)
                if current_section:
                    current_section.questions.append(current_question)

                current_question.paragraphs.append(para)
                in_material = False  # Reset material flag for normal questions
                has_material_content = False
                continue

            # Check for sub-question number (material question)
            sub_match = self.SUB_QUESTION_PATTERN.match(text)
            if sub_match and current_question:
                if is_inline_sub_section:
                    current_question.paragraphs.append(para)
                    continue

                sub_number = sub_match.group(1)

                # Mark parent as material question
                current_question.is_material_question = True

                # Create sub-question
                sub_question = QuestionBlock(f"({sub_number})", i)
                sub_question.paragraphs.append(para)
                current_question.sub_questions.append(sub_question)
                continue

            # Add paragraph to current question
            if current_question:
                if current_question.sub_questions:
                    current_question.sub_questions[-1].paragraphs.append(para)
                else:
                    current_question.paragraphs.append(para)

        # Close last question
        if current_question:
            current_question.end_index = len(self.paragraphs) - 1

        return self.sections

    def get_question_type(self, section: QuestionSection) -> str:
        """Determine question type from section name

        Args:
            section: Question section

        Returns:
            Standardized question type name
        """
        type_name = section.type_name.lower()

        if "选择" in type_name:
            return "选择题"
        elif "填空" in type_name:
            return "填空题"
        elif "解答" in type_name or "计算" in type_name:
            return "解答题"
        elif "实验" in type_name:
            return "实验题"
        elif "简答" in type_name:
            return "简答题"
        else:
            return section.type_name

    def extract_question_blocks(self) -> List[Dict]:
        """Extract all question blocks with metadata

        Returns:
            List of question dictionaries with structure info
        """
        self.parse()

        questions = []
        for section in self.sections:
            question_type = self.get_question_type(section)
            inline_sub_types = {"填空题", "简答题", "实验题", "解答题", "综合应用题"}

            for question in section.questions:
                if question_type in inline_sub_types and question.sub_questions:
                    for sub_q in question.sub_questions:
                        question.paragraphs.extend(sub_q.paragraphs)
                    question.sub_questions = []
                    question.is_material_question = False

                # Check if this is a fill-in question with sub-questions
                if question_type == "填空题" and question.sub_questions:
                    question_data = {
                        "number": question.number,
                        "type": question_type,
                        "paragraphs": question.paragraphs,
                        "is_material": True,
                        "is_fill_in": True,
                        "sub_questions": [],
                    }

                    for sub_q in question.sub_questions:
                        question_data["sub_questions"].append(
                            {
                                "number": sub_q.number,
                                "type": question_type,
                                "paragraphs": sub_q.paragraphs,
                                "parent_number": question.number,
                            }
                        )

                    questions.append(question_data)
                    continue

                question_data = {
                    "number": question.number,
                    "type": question_type,
                    "paragraphs": question.paragraphs,
                    "is_material": question.is_material_question,
                    "sub_questions": [],
                }

                # Add sub-questions if material question
                if question.is_material_question:
                    for sub_q in question.sub_questions:
                        question_data["sub_questions"].append(
                            {
                                "number": sub_q.number,
                                "type": question_type,
                                "paragraphs": sub_q.paragraphs,
                                "parent_number": question.number,
                            }
                        )

                questions.append(question_data)

        return questions
